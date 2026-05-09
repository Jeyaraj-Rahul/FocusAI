from datetime import datetime
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import HTTPException, status

from app.core.config import settings
from app.schemas.report import DocxReportRequest, ReportImage, TextAlignment
from app.utils.docx_formatting import (
    add_page_number_footer,
    add_table_of_contents,
    apply_document_defaults,
    apply_margins,
    format_paragraph,
)


class DocxReportGenerator:
    def __init__(self, output_dir: str | Path | None = None) -> None:
        self.output_dir = Path(output_dir or settings.STORAGE_DIR) / "generated"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(self, payload: DocxReportRequest) -> Path:
        document = Document()
        apply_document_defaults(document, payload.formatting)
        apply_margins(document, payload.formatting.margins)

        self._add_title_page(document, payload)
        if payload.include_table_of_contents:
            document.add_page_break()
            add_table_of_contents(document)

        document.add_page_break()
        for section_index, section in enumerate(payload.sections, start=1):
            document.add_heading(f"{section_index}. {section.heading}", level=1)
            for text in section.paragraphs:
                paragraph = document.add_paragraph(text)
                format_paragraph(paragraph, payload.formatting)

            for image_index, image in enumerate(section.images, start=1):
                self._add_image_with_caption(
                    document=document,
                    image=image,
                    section_index=section_index,
                    image_index=image_index,
                    alignment=image.alignment,
                )

        add_page_number_footer(document)
        output_path = self.output_dir / f"{self._slugify(payload.title)}-{uuid4().hex[:8]}.docx"
        document.save(output_path)
        return output_path

    def _add_title_page(self, document: Document, payload: DocxReportRequest) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(160)
        title_run = title.add_run(payload.title)
        title_run.bold = True
        title_run.font.name = payload.formatting.font_family
        title_run.font.size = Pt(payload.formatting.title_font_size)

        if payload.subtitle:
            subtitle = document.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.add_run(payload.subtitle)
            subtitle_run.font.name = payload.formatting.font_family
            subtitle_run.font.size = Pt(payload.formatting.body_font_size + 2)

        metadata_lines = [payload.author_name, payload.institution, datetime.utcnow().strftime("%d %B %Y")]
        for line in [item for item in metadata_lines if item]:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            run.font.name = payload.formatting.font_family
            run.font.size = Pt(payload.formatting.body_font_size)

    def _add_image_with_caption(
        self,
        document: Document,
        image: ReportImage,
        section_index: int,
        image_index: int,
        alignment: TextAlignment,
    ) -> None:
        image_path = Path(image.path)
        if not image_path.exists():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Image not found: {image.path}",
            )

        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = self._image_alignment(alignment)
        image_run = image_paragraph.add_run()
        image_run.add_picture(str(image_path), width=Inches(image.width_inches))

        caption_text = image.caption or f"Figure {section_index}.{image_index}"
        caption = document.add_paragraph()
        caption.alignment = self._image_alignment(alignment)
        caption_run = caption.add_run(f"Figure {section_index}.{image_index}: {caption_text}")
        caption_run.italic = True
        caption_run.font.size = Pt(10)

    def _image_alignment(self, alignment: TextAlignment):
        if alignment == TextAlignment.left:
            return WD_ALIGN_PARAGRAPH.LEFT
        if alignment == TextAlignment.right:
            return WD_ALIGN_PARAGRAPH.RIGHT
        return WD_ALIGN_PARAGRAPH.CENTER

    def _slugify(self, value: str) -> str:
        slug = "".join(character.lower() if character.isalnum() else "-" for character in value).strip("-")
        return "-".join(part for part in slug.split("-") if part)[:80] or "report"


docx_report_generator = DocxReportGenerator()
