from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from app.schemas.report import FormattingOptions, PageMargins, TextAlignment


def analyze_docx_template(path: str | Path) -> dict:
    document = Document(str(path))
    margins = _extract_margins(document)
    font_styles = _extract_font_styles(document)
    heading_styles = _extract_heading_styles(document)
    line_spacing = _extract_line_spacing(document)
    table_styles = _extract_table_styles(document)

    return {
        "source_filename": Path(path).name,
        "font_styles": font_styles,
        "margins": margins,
        "heading_styles": heading_styles,
        "line_spacing": line_spacing,
        "table_styles": table_styles,
        "formatting_options": _build_formatting_options(font_styles, margins, heading_styles, line_spacing),
    }


def formatting_options_from_profile(profile: dict) -> FormattingOptions:
    options = profile.get("formatting_options", {})
    margins = options.get("margins", {})
    return FormattingOptions(
        font_family=options.get("font_family") or "Times New Roman",
        body_font_size=options.get("body_font_size") or 12,
        heading_font_size=options.get("heading_font_size") or 16,
        title_font_size=options.get("title_font_size") or 22,
        line_spacing=options.get("line_spacing") or 1.5,
        paragraph_alignment=TextAlignment(options.get("paragraph_alignment") or TextAlignment.justify.value),
        margins=PageMargins(
            top=margins.get("top", 1.0),
            right=margins.get("right", 1.0),
            bottom=margins.get("bottom", 1.0),
            left=margins.get("left", 1.0),
        ),
    )


def _extract_margins(document: Document) -> dict:
    section = document.sections[0]
    return {
        "top": round(section.top_margin.inches, 2),
        "right": round(section.right_margin.inches, 2),
        "bottom": round(section.bottom_margin.inches, 2),
        "left": round(section.left_margin.inches, 2),
    }


def _extract_font_styles(document: Document) -> list[dict]:
    styles = []
    for style in document.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        font = style.font
        styles.append(
            {
                "name": style.name,
                "font_family": font.name,
                "font_size": round(font.size.pt) if font.size else None,
                "bold": font.bold,
                "italic": font.italic,
            }
        )
    return styles


def _extract_heading_styles(document: Document) -> list[dict]:
    headings = []
    for style in document.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.lower().startswith("heading"):
            font = style.font
            headings.append(
                {
                    "name": style.name,
                    "font_family": font.name,
                    "font_size": round(font.size.pt) if font.size else None,
                    "bold": font.bold,
                    "alignment": _alignment_name(getattr(style.paragraph_format, "alignment", None)),
                    "line_spacing": _line_spacing_value(style.paragraph_format.line_spacing),
                }
            )
    return headings


def _extract_line_spacing(document: Document) -> dict:
    values = []
    for paragraph in document.paragraphs:
        value = _line_spacing_value(paragraph.paragraph_format.line_spacing)
        if value:
            values.append(value)

    if not values:
        for style in document.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                value = _line_spacing_value(style.paragraph_format.line_spacing)
                if value:
                    values.append(value)

    most_common = Counter(values).most_common(1)
    return {
        "detected_values": values[:20],
        "primary": most_common[0][0] if most_common else 1.5,
    }


def _extract_table_styles(document: Document) -> list[dict]:
    tables = []
    for index, table in enumerate(document.tables, start=1):
        tables.append(
            {
                "index": index,
                "style": table.style.name if table.style else None,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "alignment": _alignment_name(getattr(table, "alignment", None)),
            }
        )
    return tables


def _build_formatting_options(font_styles: list[dict], margins: dict, heading_styles: list[dict], line_spacing: dict) -> dict:
    normal = next((style for style in font_styles if style["name"].lower() == "normal"), {})
    heading_one = next((style for style in heading_styles if style["name"].lower() == "heading 1"), {})

    return {
        "font_family": normal.get("font_family") or heading_one.get("font_family") or "Times New Roman",
        "body_font_size": normal.get("font_size") or 12,
        "heading_font_size": heading_one.get("font_size") or 16,
        "title_font_size": max((heading_one.get("font_size") or 16) + 6, 20),
        "line_spacing": line_spacing.get("primary") or 1.5,
        "paragraph_alignment": TextAlignment.justify.value,
        "margins": margins,
    }


def _line_spacing_value(value) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return round(float(value), 2)
    return None


def _alignment_name(value) -> str | None:
    if value is None:
        return None
    return str(value).split(".")[-1].lower()
