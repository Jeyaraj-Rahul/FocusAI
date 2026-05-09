from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from app.schemas.report import FormattingOptions, PageMargins, TextAlignment


ALIGNMENT_MAP = {
    TextAlignment.left: WD_ALIGN_PARAGRAPH.LEFT,
    TextAlignment.center: WD_ALIGN_PARAGRAPH.CENTER,
    TextAlignment.right: WD_ALIGN_PARAGRAPH.RIGHT,
    TextAlignment.justify: WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def apply_document_defaults(document: DocumentObject, formatting: FormattingOptions) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = formatting.font_family
    normal_style.font.size = Pt(formatting.body_font_size)

    for style_name, size in (("Heading 1", formatting.heading_font_size), ("Heading 2", formatting.heading_font_size - 2)):
        style = document.styles[style_name]
        style.font.name = formatting.font_family
        style.font.size = Pt(size)
        style.font.bold = True


def apply_margins(document: DocumentObject, margins: PageMargins) -> None:
    for section in document.sections:
        section.top_margin = Inches(margins.top)
        section.right_margin = Inches(margins.right)
        section.bottom_margin = Inches(margins.bottom)
        section.left_margin = Inches(margins.left)


def format_paragraph(paragraph: Paragraph, formatting: FormattingOptions, alignment: TextAlignment | None = None) -> None:
    paragraph.alignment = ALIGNMENT_MAP[alignment or formatting.paragraph_alignment]
    paragraph.paragraph_format.line_spacing = formatting.line_spacing
    paragraph.paragraph_format.space_after = Pt(8)
    for run in paragraph.runs:
        run.font.name = formatting.font_family
        run.font.size = Pt(formatting.body_font_size)


def add_field(paragraph: Paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = instruction

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def add_table_of_contents(document: DocumentObject) -> None:
    heading = document.add_heading("Table of Contents", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph()
    add_field(paragraph, r'TOC \o "1-3" \h \z \u')


def add_page_number_footer(document: DocumentObject) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run("Page ")
        add_field(paragraph, "PAGE")


def start_new_page(document: DocumentObject) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
